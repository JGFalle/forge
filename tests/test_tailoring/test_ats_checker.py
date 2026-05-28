"""Tests for pipeline/tailoring/ats_checker.py"""

import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock


def _make_doc(
    tables=0,
    text_boxes=0,
    drawings=0,
    header_text="",
    footer_text="",
    num_cols=1,
    track_changes=0,
    hyperlinks=0,
):
    """Build a minimal mock Document for ATS checker tests."""
    from docx.oxml.ns import qn

    doc = MagicMock()

    # Tables
    doc.tables = [MagicMock() for _ in range(tables)]

    # Inline shapes
    doc.inline_shapes = [MagicMock() for _ in range(0)]

    # Sections
    section = MagicMock()
    # Header
    hdr = MagicMock()
    hdr.paragraphs = [MagicMock(text=header_text)] if header_text else []
    section.header = hdr
    # Footer
    ftr = MagicMock()
    ftr.paragraphs = [MagicMock(text=footer_text)] if footer_text else []
    section.footer = ftr
    # Columns
    sectPr = MagicMock()
    cols = MagicMock()
    cols.get = MagicMock(return_value=str(num_cols))
    if num_cols > 1:
        sectPr.find = MagicMock(return_value=cols)
    else:
        sectPr.find = MagicMock(return_value=None)
    section._sectPr = sectPr
    doc.sections = [section]

    # Body XML for findall
    body = MagicMock()
    # text boxes, drawings, track changes, hyperlinks via findall
    def make_findall(text_boxes, drawings, track_changes, hyperlinks):
        def findall(xpath):
            tag = xpath.split("}")[-1] if "}" in xpath else xpath.split(":")[-1]
            if "txbxContent" in xpath:
                return [MagicMock()] * text_boxes
            if "drawing" in xpath:
                return [MagicMock()] * drawings
            if "ins" in xpath or "del" in xpath:
                return [MagicMock()] * track_changes
            if "hyperlink" in xpath:
                return [MagicMock()] * hyperlinks
            return []
        return findall

    body.findall = make_findall(text_boxes, drawings, track_changes, hyperlinks)
    doc.element = MagicMock()
    doc.element.body = body
    return doc


# ── Test with real DOCX ───────────────────────────────────────────────────────

def test_clean_resume_passes():
    """The actual base resume (if present) should score 100/100."""
    path = Path("assets/base_resume_v9.docx")
    if not path.exists():
        pytest.skip("Base resume not available in test environment")
    from pipeline.tailoring.ats_checker import check
    result = check(path)
    assert result["score"] == 100
    assert result["grade"] == "PASS"
    assert result["findings"] == []


# ── Unit tests via mocks ──────────────────────────────────────────────────────

def test_table_detected():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        mock_doc_cls.return_value = _make_doc(tables=1)
        result = check(Path("fake.docx"))
    issues = [f["issue"] for f in result["findings"]]
    assert "tables" in issues
    assert result["score"] < 100
    assert result["grade"] != "PASS"


def test_text_box_detected():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        mock_doc_cls.return_value = _make_doc(text_boxes=1)
        result = check(Path("fake.docx"))
    issues = [f["issue"] for f in result["findings"]]
    assert "text_boxes" in issues


def test_header_content_detected():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        mock_doc_cls.return_value = _make_doc(header_text="Jack Falle | Atlanta, GA")
        result = check(Path("fake.docx"))
    issues = [f["issue"] for f in result["findings"]]
    assert "header_content" in issues


def test_clean_doc_passes():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        mock_doc_cls.return_value = _make_doc()
        result = check(Path("fake.docx"))
    assert result["score"] == 100
    assert result["grade"] == "PASS"
    assert result["findings"] == []


def test_score_floor_is_zero():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        mock_doc_cls.return_value = _make_doc(
            tables=5, text_boxes=3, header_text="name here",
            num_cols=2, track_changes=2
        )
        result = check(Path("fake.docx"))
    assert result["score"] >= 0


def test_grade_fail_threshold():
    from pipeline.tailoring.ats_checker import check
    with patch("pipeline.tailoring.ats_checker.Document") as mock_doc_cls:
        # tables + text boxes = -25 -25 = 50 → FAIL
        mock_doc_cls.return_value = _make_doc(tables=1, text_boxes=1)
        result = check(Path("fake.docx"))
    assert result["grade"] == "FAIL"
