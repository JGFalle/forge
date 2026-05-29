"""ATS compatibility checker for tailored resume DOCX output.

Scans the document for elements that Workday, Greenhouse, Lever, and iCIMS
parsers cannot reliably read. Run on the output DOCX after tailoring.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from utils.logging import get_logger

logger = get_logger(__name__)

_SEVERITY_PENALTY = {
    "high": 25,
    "medium": 12,
    "low": 5,
}


def check(doc_path: Path) -> dict:
    """
    Analyze doc_path for ATS-unfriendly elements.
    Returns {score, grade, findings, doc_path}.
    """
    doc = Document(doc_path)
    body = doc.element.body
    findings = []

    # 1. Tables: layout tables cause ATS parsers to skip or scramble content
    table_count = len(doc.tables)
    if table_count > 0:
        findings.append({
            "issue": "tables",
            "count": table_count,
            "severity": "high",
            "detail": (
                f"{table_count} table(s) found. Most ATS parsers skip table cells "
                "or read them out of sequence, corrupting the resume content."
            ),
            "fix": "Reformat using tab stops or plain paragraphs instead of tables.",
        })

    # 2. Text boxes: w:txbxContent is invisible to most ATS
    text_boxes = body.findall(f".//{qn('w:txbxContent')}")
    if text_boxes:
        findings.append({
            "issue": "text_boxes",
            "count": len(text_boxes),
            "severity": "high",
            "detail": (
                f"{len(text_boxes)} text box(es) found. Text inside drawing shapes "
                "is completely invisible to ATS parsers."
            ),
            "fix": "Move all text out of text boxes into body paragraphs.",
        })

    # 3. Drawings / images (w:drawing covers both inline images and shapes)
    drawings = body.findall(f".//{qn('w:drawing')}")
    # Subtract drawings that are part of text boxes already counted
    image_count = len(drawings) - len(text_boxes)
    if image_count > 0:
        findings.append({
            "issue": "images_or_shapes",
            "count": image_count,
            "severity": "medium",
            "detail": (
                f"{image_count} image or drawing element(s). ATS parsers ignore images. "
                "Icons used as section dividers are usually harmless; logos or infographics are not."
            ),
            "fix": "Remove decorative images. Replace icon bullets with Unicode characters in body text.",
        })

    # 4. Header content. ATS often skips headers/footers entirely
    has_header_content = False
    has_footer_content = False
    for section in doc.sections:
        if not has_header_content:
            try:
                hdr = section.header
                if hdr and any(p.text.strip() for p in hdr.paragraphs):
                    has_header_content = True
            except Exception:
                pass
        if not has_footer_content:
            try:
                ftr = section.footer
                if ftr and any(p.text.strip() for p in ftr.paragraphs):
                    has_footer_content = True
            except Exception:
                pass

    if has_header_content:
        findings.append({
            "issue": "header_content",
            "count": 1,
            "severity": "high",
            "detail": (
                "The document header contains text. If contact information lives here "
                "(name, email, phone), most ATS parsers will not capture it."
            ),
            "fix": "Move name and contact info into the body of the document.",
        })

    if has_footer_content:
        findings.append({
            "issue": "footer_content",
            "count": 1,
            "severity": "low",
            "detail": "Footer contains text (likely page numbers). Low risk for most ATS.",
            "fix": "No action required unless footer contains substantive content.",
        })

    # 5. Multi-column layout
    for section in doc.sections:
        try:
            sectPr = section._sectPr
            cols = sectPr.find(qn("w:cols"))
            if cols is not None:
                num_cols = int(cols.get(qn("w:num"), 1))
                if num_cols > 1:
                    findings.append({
                        "issue": "multi_column_layout",
                        "count": num_cols,
                        "severity": "high",
                        "detail": (
                            f"{num_cols}-column layout detected. ATS parsers linearize columns "
                            "top-to-bottom, left-to-right — mixing experience from separate columns."
                        ),
                        "fix": "Convert to a single-column layout.",
                    })
                    break
        except Exception:
            pass

    # 6. Track changes: unaccepted changes leave ghost text that confuses parsers
    insertions = body.findall(f".//{qn('w:ins')}")
    deletions = body.findall(f".//{qn('w:del')}")
    if insertions or deletions:
        findings.append({
            "issue": "track_changes",
            "count": len(insertions) + len(deletions),
            "severity": "medium",
            "detail": (
                f"{len(insertions) + len(deletions)} tracked change(s) present. "
                "Deleted text may still be parsed; inserted text may be duplicated."
            ),
            "fix": "Accept all changes before saving: Review > Accept All Changes.",
        })

    # 7. Hyperlinks: flag if excessive (LinkedIn/email are expected, lots of URLs are noise)
    hyperlinks = body.findall(f".//{qn('w:hyperlink')}")
    if len(hyperlinks) > 6:
        findings.append({
            "issue": "excessive_hyperlinks",
            "count": len(hyperlinks),
            "severity": "low",
            "detail": (
                f"{len(hyperlinks)} hyperlinks found. Most ATS strip URLs. "
                "Keep only contact-info links; remove embedded URLs from bullet text."
            ),
            "fix": "Limit to name/email/LinkedIn in the header block.",
        })

    # Score: start at 100, subtract weighted penalties (capped per issue)
    score = 100
    for f in findings:
        penalty = _SEVERITY_PENALTY.get(f["severity"], 0)
        score -= penalty
    score = max(score, 0)

    grade = "PASS" if score >= 85 else ("WARN" if score >= 65 else "FAIL")

    logger.info(
        "ATS check: %s — score %d/100 [%s], %d finding(s)",
        doc_path.name, score, grade, len(findings)
    )
    return {
        "score": score,
        "grade": grade,
        "findings": findings,
        "doc_path": str(doc_path),
    }


def display_report(result: dict) -> None:
    """Print ATS report to terminal."""
    score = result["score"]
    grade = result["grade"]
    findings = result["findings"]
    width = 60

    grade_display = {"PASS": "PASS", "WARN": "WARN", "FAIL": "FAIL"}.get(grade, grade)
    print(f"\n{'='*width}")
    print(f"  ATS COMPATIBILITY   Score: {score}/100  [{grade_display}]")
    print(f"{'='*width}")

    if not findings:
        print("  No issues — document is ATS-friendly.\n")
        return

    severity_rank = {"high": 0, "medium": 1, "low": 2}
    for f in sorted(findings, key=lambda x: severity_rank.get(x["severity"], 3)):
        icon = {"high": "✗", "medium": "!", "low": "~"}.get(f["severity"], "?")
        label = f["issue"].replace("_", " ").upper()
        print(f"\n  [{icon}] {label}  ({f['severity'].upper()})")
        print(f"      {f['detail']}")
        print(f"      Fix: {f['fix']}")
    print()


def save_report(result: dict, output_dir: Path, slug: str) -> Path:
    """Save ATS report as plain text alongside other research artifacts."""
    output_dir.mkdir(parents=True, exist_ok=True)
    out = output_dir / f"ats_check_{slug}.txt"

    lines = [
        f"ATS COMPATIBILITY REPORT — {slug}",
        f"Score: {result['score']}/100  [{result['grade']}]",
        f"Document: {result['doc_path']}",
        "",
    ]
    if not result["findings"]:
        lines.append("No issues found — document is ATS-friendly.")
    else:
        severity_rank = {"high": 0, "medium": 1, "low": 2}
        for f in sorted(result["findings"], key=lambda x: severity_rank.get(x["severity"], 3)):
            lines += [
                f"[{f['severity'].upper()}] {f['issue'].replace('_', ' ').upper()}  (count: {f['count']})",
                f"  {f['detail']}",
                f"  Fix: {f['fix']}",
                "",
            ]
    out.write_text("\n".join(lines), encoding="utf-8")
    return out
