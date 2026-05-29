from __future__ import annotations

import copy
import logging
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


def _load_role_anchors() -> dict[str, str]:
    """Build role_id -> anchor_text map from config career_history."""
    from utils.config import get
    history = get("career_history", [])
    anchors = {}
    for role in history:
        role_id = role.get("id")
        anchor = role.get("anchor_text")
        if role_id and anchor:
            anchors[role_id] = anchor
    return anchors


class DocxModifier:
    @property
    def ROLE_ANCHORS(self) -> dict[str, str]:
        if not hasattr(self, "_role_anchors_cache"):
            self._role_anchors_cache = _load_role_anchors()
        return self._role_anchors_cache

    def __init__(self, doc: Document, tailoring: dict, dry_run: bool = False):
        self.doc = doc
        self.t = tailoring
        self.dry_run = dry_run
        self.paragraphs = self.doc.paragraphs

    def apply_all(self) -> None:
        logger.info("DocxModifier: apply_all (dry_run=%s)", self.dry_run)
        if "headline" in self.t:
            self._apply_headline()
        if "summary" in self.t:
            self._apply_summary()
        if "competencies" in self.t:
            self._apply_competencies()
        if "experience_modifications" in self.t:
            for mod in self.t["experience_modifications"]:
                self._apply_experience_mod(mod)
        if "technical_skills" in self.t:
            self._apply_technical_skills()
        logger.info("DocxModifier: apply_all complete.")

    # Paragraph utilities

    def _find_paragraph_by_text(self, search_text: str, exact: bool = False) -> Optional[int]:
        target = search_text.upper()
        for i, p in enumerate(self.paragraphs):
            text = p.text.strip().upper()
            if exact:
                if text == target:
                    return i
            else:
                if target in text:
                    return i
        return None

    def _set_paragraph_text(self, paragraph: Paragraph, new_text: str) -> None:
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    # Section appliers

    def _apply_headline(self) -> None:
        new_headline = self.t["headline"]
        idx = self._find_paragraph_by_text("STRATEGIC OPERATIONS")
        if idx is None:
            idx = self._find_paragraph_by_text("LEADER")
        if idx is None:
            logger.warning("Headline paragraph not found; skipping.")
            return
        p = self.paragraphs[idx]
        logger.info("Updating headline at [%d]: '%s'", idx, p.text.strip()[:80])
        if self.dry_run:
            logger.info("DRY RUN: headline -> %s", new_headline)
            return
        self._set_paragraph_text(p, new_headline)

    def _apply_summary(self) -> None:
        new_summary = self.t["summary"]
        heading_idx = self._find_paragraph_by_text("PROFESSIONAL SUMMARY")
        if heading_idx is None:
            logger.warning("PROFESSIONAL SUMMARY heading not found; skipping.")
            return
        idx = None
        for i in range(heading_idx + 1, len(self.paragraphs)):
            if self.paragraphs[i].text.strip():
                idx = i
                break
        if idx is None:
            logger.warning("Summary content paragraph not found; skipping.")
            return
        if self.dry_run:
            logger.info("DRY RUN: summary -> %d chars", len(new_summary))
            return
        self._set_paragraph_text(self.paragraphs[idx], new_summary)

    def _apply_competencies(self) -> None:
        comps = self.t.get("competencies", [])
        if not comps:
            return
        heading_idx = self._find_paragraph_by_text("CORE COMPETENCIES")
        if heading_idx is None:
            logger.warning("CORE COMPETENCIES heading not found; skipping.")
            return

        comp_start = heading_idx + 1
        stop_markers = ["PROFESSIONAL EXPERIENCE", "PROFESSIONAL SUMMARY", "EDUCATION", "TECHNICAL SKILLS"]
        comp_end = comp_start
        for i in range(comp_start, len(self.paragraphs)):
            text = self.paragraphs[i].text.strip()
            if not text:
                comp_end = i
                break
            if any(m.upper() in text.upper() for m in stop_markers):
                comp_end = i
                break
            comp_end = i + 1

        num_rows = comp_end - comp_start
        logger.info("CORE COMPETENCIES: %d rows, %d provided.", num_rows, len(comps))

        if self.dry_run:
            for i, c in enumerate(comps):
                logger.info("  DRY RUN comp %d: %s", i + 1, c)
            return

        if len(comps) == num_rows:
            for i, new_text in enumerate(comps):
                idx = comp_start + i
                p = self.paragraphs[idx]
                runs = p.runs
                if not runs:
                    continue
                accumulated = ""
                has_label = False
                for run in runs:
                    accumulated += run.text
                    if ":" in accumulated:
                        has_label = True
                        break
                if has_label:
                    label_text = accumulated.split(":")[0] + ": "
                    for old_run in p._element.findall(qn("w:r")):
                        p._element.remove(old_run)
                    for old_hlink in p._element.findall(qn("w:hyperlink")):
                        p._element.remove(old_hlink)
                    label_run_el = p._element.makeelement(qn("w:r"), {})
                    rPr = runs[0]._element.find(qn("w:rPr"))
                    if rPr is not None:
                        label_run_el.insert(0, copy.deepcopy(rPr))
                    label_t = label_run_el.makeelement(qn("w:t"), {})
                    label_t.text = label_text
                    label_t.set(qn("xml:space"), "preserve")
                    label_run_el.append(label_t)
                    p._element.append(label_run_el)
                    content_run_el = p._element.makeelement(qn("w:r"), {})
                    non_bold_run = next((r for r in runs if not r.bold), None)
                    if non_bold_run is not None and non_bold_run._element.find(qn("w:rPr")) is not None:
                        content_rPr = copy.deepcopy(non_bold_run._element.find(qn("w:rPr")))
                    else:
                        content_rPr = copy.deepcopy(runs[0]._element.find(qn("w:rPr")))
                        for tag in [qn("w:b"), qn("w:bCs")]:
                            el = content_rPr.find(tag)
                            if el is not None:
                                content_rPr.remove(el)
                    content_run_el.insert(0, content_rPr)
                    content_t = content_run_el.makeelement(qn("w:t"), {})
                    content_t.text = new_text
                    content_t.set(qn("xml:space"), "preserve")
                    content_run_el.append(content_t)
                    p._element.append(content_run_el)
                    logger.info("  Comp [%d]: '%s' -> '%s'", idx, label_text.strip(), new_text[:60])
                else:
                    self._set_paragraph_text(p, new_text)
        else:
            self._set_paragraph_text(self.paragraphs[comp_start], " • ".join(comps))
            logger.info("  Comp count mismatch; wrote single line.")

    # Experience

    def _find_role_header_index(self, role_id: str) -> Optional[int]:
        anchor = self.ROLE_ANCHORS.get(role_id)
        if not anchor:
            logger.warning("Unknown role_identifier: %s", role_id)
            return None
        idx = self._find_paragraph_by_text(anchor)
        if idx is None:
            logger.warning("Role header not found for %s (anchor='%s')", role_id, anchor)
        return idx

    def _find_role_bullet_range(self, header_idx: int) -> tuple[int, int]:
        """Return (start, end) indices of bullet paragraphs (those with w:numPr)."""
        stop_markers = [
            "PROFESSIONAL EXPERIENCE", "EDUCATION", "TECHNICAL SKILLS",
            "CORE COMPETENCIES", "PROFESSIONAL SUMMARY",
        ]
        stop_markers.extend(self.ROLE_ANCHORS.values())

        start = None
        end = None
        for i in range(header_idx + 1, len(self.paragraphs)):
            text = self.paragraphs[i].text.strip()
            if any(m.upper() in text.upper() for m in stop_markers) and i != header_idx:
                break
            pPr = self.paragraphs[i]._element.find(qn("w:pPr"))
            has_numPr = pPr is not None and pPr.find(qn("w:numPr")) is not None
            if has_numPr:
                if start is None:
                    start = i
                end = i + 1

        if start is None:
            return header_idx + 1, header_idx + 1
        return start, end

    def _apply_experience_mod(self, mod: dict) -> None:
        role_id = mod.get("role_identifier")
        new_bullets = mod.get("replacement_bullets", [])
        new_lead_in = mod.get("replacement_lead_in")

        if not new_bullets and not new_lead_in:
            return

        header_idx = self._find_role_header_index(role_id)
        if header_idx is None:
            return

        if new_lead_in:
            for i in range(header_idx + 1, len(self.paragraphs)):
                text = self.paragraphs[i].text.strip()
                if not text:
                    continue
                pPr = self.paragraphs[i]._element.find(qn("w:pPr"))
                has_numPr = pPr is not None and pPr.find(qn("w:numPr")) is not None
                if not has_numPr:
                    if self.dry_run:
                        logger.info("DRY RUN: lead-in for %s -> %s", role_id, new_lead_in[:80])
                    else:
                        self._set_paragraph_text(self.paragraphs[i], new_lead_in)
                        logger.info("Replaced lead-in for '%s' at [%d].", role_id, i)
                    break
                else:
                    logger.warning("No lead-in paragraph found for %s before bullets.", role_id)
                    break

        if not new_bullets:
            return

        start, end = self._find_role_bullet_range(header_idx)
        old_count = end - start
        logger.info("Role '%s': bullets [%d:%d] (%d existing) -> %d new.", role_id, start, end, old_count, len(new_bullets))

        if self.dry_run:
            for i, b in enumerate(new_bullets):
                logger.info("  DRY RUN bullet %d: %s", i + 1, b[:100])
            return

        if start >= end:
            logger.warning("No bullets found for '%s'; skipping.", role_id)
            return

        template_element = copy.deepcopy(self.paragraphs[start]._element)
        for i in range(end - 1, start - 1, -1):
            el = self.paragraphs[i]._element
            el.getparent().remove(el)

        insert_after = self.paragraphs[start - 1]._element
        for bullet_text in new_bullets:
            new_p = copy.deepcopy(template_element)
            for old_run in new_p.findall(qn("w:r")):
                new_p.remove(old_run)
            for old_hlink in new_p.findall(qn("w:hyperlink")):
                new_p.remove(old_hlink)

            # Strip Word list numbering (w:numPr) so Workday autofill reads plain text
            # instead of resetting to "1." on every bullet. Visual indentation is
            # preserved via the paragraph indent already set on the template style.
            pPr = new_p.find(qn("w:pPr"))
            if pPr is not None:
                num_pr = pPr.find(qn("w:numPr"))
                if num_pr is not None:
                    pPr.remove(num_pr)

            new_r = new_p.makeelement(qn("w:r"), {})
            template_runs = template_element.findall(qn("w:r"))
            if template_runs:
                template_rPr = template_runs[0].find(qn("w:rPr"))
                if template_rPr is not None:
                    new_r.insert(0, copy.deepcopy(template_rPr))
            new_t = new_r.makeelement(qn("w:t"), {})
            # Prefix with bullet character so the point is visually clear without Word numbering
            new_t.text = f"•  {bullet_text}" if not bullet_text.startswith("•") else bullet_text
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            new_p.append(new_r)
            insert_after.addnext(new_p)
            insert_after = new_p

        self.paragraphs = self.doc.paragraphs
        logger.info("Role '%s': replaced %d bullets with %d.", role_id, old_count, len(new_bullets))

    def _apply_technical_skills(self) -> None:
        skills = self.t.get("technical_skills", [])
        if not skills:
            return
        heading_idx = self._find_paragraph_by_text("TECHNICAL SKILLS")
        if heading_idx is None:
            logger.warning("TECHNICAL SKILLS heading not found; skipping.")
            return
        skill_start = heading_idx + 1

        if self.dry_run:
            for i, s in enumerate(skills):
                logger.info("  DRY RUN skill row %d: %s", i + 1, s)
            return

        for i, new_text in enumerate(skills):
            idx = skill_start + i
            if idx >= len(self.paragraphs):
                logger.warning("Not enough paragraphs for skill row %d.", i + 1)
                break
            p = self.paragraphs[idx]
            runs = p.runs
            if not runs:
                continue
            accumulated = ""
            for j, run in enumerate(runs):
                accumulated += run.text
                if ":" in accumulated:
                    break
            label_text = accumulated.split(":")[0] + ": "
            for old_run in p._element.findall(qn("w:r")):
                p._element.remove(old_run)
            for old_hlink in p._element.findall(qn("w:hyperlink")):
                p._element.remove(old_hlink)
            label_run_el = p._element.makeelement(qn("w:r"), {})
            rPr = runs[0]._element.find(qn("w:rPr"))
            if rPr is not None:
                label_run_el.insert(0, copy.deepcopy(rPr))
            label_t = label_run_el.makeelement(qn("w:t"), {})
            label_t.text = label_text
            label_t.set(qn("xml:space"), "preserve")
            label_run_el.append(label_t)
            p._element.append(label_run_el)
            content_run_el = p._element.makeelement(qn("w:r"), {})
            non_bold_run = next((r for r in runs if not r.bold), None)
            if non_bold_run is not None and non_bold_run._element.find(qn("w:rPr")) is not None:
                content_rPr = copy.deepcopy(non_bold_run._element.find(qn("w:rPr")))
            else:
                content_rPr = copy.deepcopy(runs[0]._element.find(qn("w:rPr")))
                for tag in [qn("w:b"), qn("w:bCs")]:
                    el = content_rPr.find(tag)
                    if el is not None:
                        content_rPr.remove(el)
            content_run_el.insert(0, content_rPr)
            content_t = content_run_el.makeelement(qn("w:t"), {})
            content_t.text = new_text
            content_t.set(qn("xml:space"), "preserve")
            content_run_el.append(content_t)
            p._element.append(content_run_el)
            logger.info("  Skill row [%d]: '%s' -> '%s'", idx, label_text.strip(), new_text[:60])
