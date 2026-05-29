"""Generate a formatted cover letter DOCX from a tailoring JSON."""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from utils.config import get
from utils.text import strip_em_dashes, word_count

logger = logging.getLogger(__name__)

_NAVY = RGBColor(0x0E, 0x36, 0x89)
_GRAY = RGBColor(0x59, 0x59, 0x59)


def _cl_filename() -> str:
    resume_fn = get("person.resume_filename", "YourName_Resume.docx")
    name_part = resume_fn.replace("_Resume.docx", "")
    return f"{name_part}_Cover_Letter"


def run_txt(tailoring_json_path: Path, output_dir: Path) -> Path:
    """Plain text version: for pasting into ATS text boxes."""
    with open(tailoring_json_path, encoding="utf-8") as f:
        data = json.load(f)
    body = strip_em_dashes(data.get("cover_letter", "").strip())
    if not body:
        raise ValueError(f"No 'cover_letter' field in {tailoring_json_path}")

    name = get("person.name", "")
    phone = get("person.phone", "")
    email = get("person.email", "")
    linkedin = get("person.linkedin", "")
    company = data.get("company", "")
    role = data.get("role", "")
    salutation = data.get("cover_letter_salutation", "Hiring Team")

    lines = [
        name,
        f"{phone} | {email} | {linkedin}",
        "",
        datetime.now().strftime("%B %d, %Y"),
        "",
        f"Dear {salutation},",
        "",
    ]
    if company and role:
        lines += [f"Re: {role} — {company}", ""]
    lines += [body, "", "Sincerely,", name]

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"{_cl_filename()}.txt"
    out_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("Cover letter (txt) saved: %s", out_path)
    return out_path


def run(tailoring_json_path: Path, output_dir: Path) -> Path:
    with open(tailoring_json_path, encoding="utf-8") as f:
        data = json.load(f)

    body = data.get("cover_letter", "").strip()
    if not body:
        raise ValueError(f"No 'cover_letter' field in {tailoring_json_path}")

    body = strip_em_dashes(body)
    wc = word_count(body)
    max_words = get("tailoring.cover_letter_words_max", 150)
    if wc > max_words:
        raise ValueError(
            f"Cover letter is {wc} words — exceeds {max_words}-word limit. "
            f"Edit the cover_letter field in the tailoring JSON and re-run."
        )

    company = data.get("company", "Company")
    role = data.get("role", "Role")
    salutation = data.get("cover_letter_salutation", "Hiring Team")

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"{_cl_filename()}.docx"

    _build_docx(body, company, role, salutation, out_path)
    logger.info("Cover letter saved: %s (%d words)", out_path, wc)
    return out_path


def _build_docx(body: str, company: str, role: str, salutation: str, out_path: Path) -> None:
    name = get("person.name", "")
    phone = get("person.phone", "")
    email = get("person.email", "")
    linkedin = get("person.linkedin", "")
    location = get("person.location", "")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Calibri"
    r.font.color.rgb = _NAVY

    p = doc.add_paragraph()
    r = p.add_run(f"{location} | {phone} | {email} | {linkedin}")
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = _GRAY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("_" * 72)
    r.font.size = Pt(2)
    r.font.color.rgb = _NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(datetime.now().strftime("%B %d, %Y"))
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"Dear {salutation},")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"Re: {role} — {company}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(body)
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Sincerely,")
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    doc.save(str(out_path))
